import re
import requests
import json

import sys

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# doc.add_heading('The REAL meaning of the universe')
# table = doc.add_table(6, 4, style="Table Grid")
# table.cell(0, 1).merge(table.cell(0, 3))
# table.cell(0, 0).text = 'URL'
# table.cell(0, 1).text = "http://www.baidu.com"
# table.cell(1, 1).merge(table.cell(1, 3))
# table.cell(1, 0).text = "请求方式"
# table.cell(1, 1).text = "POST"
# table.cell(2, 0).text = "请求参数"
# table.cell(2, 1).text = "参数名"
# table.cell(2, 2).text = "数据类型"
# table.cell(2, 3).text = "是否必填"


r = requests.get("http://127.0.0.1:8081/v2/api-docs")
text = str(r.content, encoding="utf-8")
text_json = json.loads(text)


for item in text_json['paths']:
    head = ""
    url = ""
    request_method = ""
    params_list = []
    # URL
    url = item

    # 请求方式
    value = text_json['paths'][item]
    for httpMethod in value:
        request_method = httpMethod.upper()
        content = value[httpMethod]
        head = content['summary']
        if 'parameters' in content:
            params = content['parameters']
            params_list = params

    # 写入表格
    doc.add_heading(head)
    table = doc.add_table(3+len(params_list), 4, style="Table Grid")
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 0).text = 'URL'
    table.cell(0, 1).text = url
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 0).text = "请求方式"
    table.cell(1, 1).text = request_method

    table.cell(2, 0).text = "请求参数名"
    table.cell(2, 1).text = "参数名"
    table.cell(2, 2).text = "数据类型"
    table.cell(2, 3).text = "是否必填"
    row = 3
    for param in params_list:
        print(param)
        if 'description' in param:
            table.cell(row, 0).text = param['description']

        if 'name' in param:
            table.cell(row, 1).text = param['name']
        if 'schema' in param:
            temp = param['schema']
            if 'type' in temp:
                table.cell(row, 2).text = param['schema']['type']
        if 'required' in param:
            table.cell(row, 3).text = str(param['required'])
        row += 1

doc.save("demo.docx")
